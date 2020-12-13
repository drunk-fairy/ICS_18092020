import docx

from data import data_

head = data_[0]

doc = docx.Document()

doc.add_heading('Аналіз руху основних засобів', 1)

datatable = doc.add_table(rows = 1, cols = 7)
datatable.style = 'Table Grid'
hdr_Cells = datatable.rows[0].cells

for i in range (0, len(head)):
    hdr_Cells[i].text = head[i]

for pidpr, vid, zal18, nad, vib, zal19, zmin in data_:
    row_Cells = datatable.add_row().cells
    row_Cells[0].text = pidpr
    row_Cells[1].text = str(vid)
    row_Cells[2].text = str(zal18)
    row_Cells[3].text = str(nad)
    row_Cells[4].text = str(vib)
    row_Cells[5].text = str(zal19)
    row_Cells[6].text = str(zmin)

def removeRow(datatable, row):
    tbl = datatable._tbl
    tr = row._tr
    tbl.remove(tr)

row = datatable.rows[0]
removeRow(datatable, row)

def saveDOCX():
    doc.save('e:/University/ВКН індивідуальний_проект/table.docx')
    # Save the file to C:/Users/Admin:
    #doc.save('table.docx')